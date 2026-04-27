import docx
from utils.logger import setup_logger

logger = setup_logger("NativeDOCX")

class DocxProcessor:
    """
    Extracts text and embedded images from DOCX files.
    """

    def __init__(self):
        self.extracted_data = []

    def extract_text_with_metadata(self, docx_path):
        """
        Extracts paragraphs and detects basic styling for segmentation.
        """
        logger.info(f"Starting native extraction for: {docx_path}")
        try:
            doc = docx.Document(docx_path)
            full_content = []

            for i, para in enumerate(doc.paragraphs):
                if not para.text.strip():
                    continue

                # In DOCX, styling is often in the first 'run' of a paragraph
                # We'll take the metadata from the first run as a representative
                first_run = para.runs[0] if para.runs else None
                
                metadata = {
                    "text": para.text.strip(),
                    "font_size": first_run.font.size.pt if first_run and first_run.font.size else 12.0,
                    "is_bold": first_run.bold if first_run else False,
                    "style": para.style.name, # Paragraph styles like 'Heading 1'
                    "para_index": i
                }
                full_content.append(metadata)

            logger.info(f"Successfully extracted {len(full_content)} paragraphs.")
            return full_content

        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}", exc_info=True)
            return []

    def extract_images(self, docx_path, output_folder):
        """
        Extracts images from the Word document's internal XML structure.
        """
        logger.info("Extracting images from DOCX...")
        try:
            doc = docx.Document(docx_path)
            image_count = 0

            for rel in doc.part.rels.values():
                # Check if the relation target is an image
                if "image" in rel.target_ref:
                    image_count += 1
                    image_data = rel.target_part.blob
                    # Get extension from the target reference (e.g., 'media/image1.png')
                    ext = rel.target_ref.split('.')[-1]
                    
                    image_name = f"docx_img_{image_count}.{ext}"
                    image_path = f"{output_folder}/{image_name}"

                    with open(image_path, "wb") as f:
                        f.write(image_data)

            logger.info(f"Extracted {image_count} images from DOCX.")
            return image_count
        except Exception as e:
            logger.error(f"DOCX image extraction failed: {e}")
            return 0