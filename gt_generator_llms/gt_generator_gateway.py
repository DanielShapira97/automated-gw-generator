"""
Ground-truth pipeline using the Jeenai completion HTTP API (same contract as
completion-service.client.ts: minimal JSON body, messages with type "message",
response field outputText).

Default URL: https://completion.stg.jeenai.app/api/v1/completions (override with
COMPLETIONS_API_URL). Optional: COMPLETIONS_MODEL, COMPLETIONS_BEARER_TOKEN.
For the FastAPI webapp, set COMPLETIONS_USE_GATEWAY=1 or COMPLETIONS_API_URL.
"""
import argparse
import base64
import os
from typing import Any

import docx
import fitz
from dotenv import load_dotenv
from pdf2image import convert_from_path

from .http_completions_client import complete_messages

load_dotenv()

VISION_PROMPT = (
    "Extract all text from this image exactly. No data loss. Format any tables using | and - "
    "characters to draw them. Do NOT extract text from pictures, diagrams, or charts. "
    "Do NOT describe images or pictures. No markdown formatting other than for tables."
)

BLOCK_PROMPT = (
    "Divide the text into LARGE logical blocks using ==== markers.\n"
    "1. DO NOT DELETE ANY TEXT. NO SUMMARIZING.\n"
    "2. NO MARKDOWN symbols except for tables using | and -.\n"
    "3. Every block must start and end with ====.\n"
    "4. Return ONLY the blocks.\n\n"
    "Text:\n\n"
)


def extract_images_from_pdf(pdf_path, doc_folder):
    doc = fitz.open(pdf_path)
    image_count = 0
    for page_index in range(len(doc)):
        for _img_index, img in enumerate(doc.get_page_images(page_index)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]

            image_count += 1
            img_filename = f"image_{image_count}.{image_ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(image_bytes)
    print(f"Extracted {image_count} images from PDF.")


def extract_images_from_docx(docx_path, doc_folder):
    doc = docx.Document(docx_path)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            img_data = rel.target_part.blob
            ext = rel.target_ref.split(".")[-1]
            img_filename = f"image_{image_count}.{ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(img_data)
    print(f"Extracted {image_count} images from DOCX.")


def _messages_vision_jpeg(jpeg_path: str) -> list[dict[str, Any]]:
    with open(jpeg_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": VISION_PROMPT},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:image/jpeg;base64,{b64}"},
                },
            ],
        }
    ]


def get_raw_text_content(input_path, doc_folder, poppler_path):
    file_ext = os.path.splitext(input_path)[1].lower()
    raw_content = ""

    if file_ext == ".docx":
        print(f"Reading DOCX: {input_path}")
        doc = docx.Document(input_path)
        for p in doc.paragraphs:
            raw_content += p.text + "\n"
        for t in doc.tables:
            for r in t.rows:
                raw_content += " | ".join([c.text.strip() for c in r.cells]) + "\n"
        extract_images_from_docx(input_path, doc_folder)

    elif file_ext == ".pdf":
        extract_images_from_pdf(input_path, doc_folder)
        doc = fitz.open(input_path)
        print(f"Total pages to process: {len(doc)}")

        for i in range(len(doc)):
            print(f"--- Processing Page {i+1} ---")
            page = doc.load_page(i)
            pix = page.get_pixmap(dpi=200)
            temp_path = os.path.join(doc_folder, f"p_{i}.jpg")
            pix.save(temp_path)

            try:
                text = complete_messages(_messages_vision_jpeg(temp_path))
                if text:
                    raw_content += text + "\n\n"
            except Exception as e:
                print(f"Error on Page {i+1}: {e}")

            os.remove(temp_path)

    return raw_content


def get_raw_text_content_poppler(input_path, doc_folder, poppler_path):
    """PDF via pdf2image + gateway vision (same as gt_generator_llm flow)."""
    raw_content = ""
    extract_images_from_pdf(input_path, doc_folder)
    pages = convert_from_path(input_path, poppler_path=poppler_path)
    for i, page in enumerate(pages):
        print(f"Vision Processing: Page {i+1}...")
        temp_path = os.path.join(doc_folder, f"temp_p{i}.jpg")
        page.save(temp_path, "JPEG")
        try:
            text = complete_messages(_messages_vision_jpeg(temp_path))
            raw_content += text + "\n\n"
        finally:
            if os.path.isfile(temp_path):
                os.remove(temp_path)
    return raw_content


def structure_into_semantic_blocks(raw_text):
    if not raw_text.strip():
        return "Error: No text extracted."
    print("Organizing text into logical blocks...")
    messages = [{"role": "user", "content": f"{BLOCK_PROMPT}{raw_text}"}]
    try:
        return complete_messages(messages).strip()
    except Exception as e:
        print(f"Structuring error: {e}")
        return raw_text


def process_document(input_path, output_base="evaluation_results", poppler=None):
    if not input_path:
        print("Missing input_path")
        return None

    doc_name = os.path.splitext(os.path.basename(input_path))[0]
    doc_folder = os.path.join(output_base, doc_name)
    os.makedirs(doc_folder, exist_ok=True)

    file_ext = os.path.splitext(input_path)[1].lower()
    if file_ext == ".pdf" and os.getenv("USE_POPPLER_FOR_PDF", "").strip().lower() in (
        "1",
        "true",
        "yes",
    ):
        raw_text = get_raw_text_content_poppler(input_path, doc_folder, poppler)
    else:
        raw_text = get_raw_text_content(input_path, doc_folder, poppler)

    final_text = structure_into_semantic_blocks(raw_text)

    out_file = os.path.join(doc_folder, f"{doc_name} ground truth.txt")
    with open(out_file, "w", encoding="utf-8") as f:
        f.write(final_text)

    print(f"\n--- SUCCESS! Created: {out_file} ---")
    return out_file


def main():
    # .env is optional; load_dotenv() does nothing if the file is absent.
    parser = argparse.ArgumentParser(
        description="Generate ground truth via COMPLETIONS_API_URL (no .env required)."
    )
    parser.add_argument(
        "input_file",
        nargs="?",
        help="Path to a PDF or DOCX file",
    )
    parser.add_argument(
        "-o",
        "--output",
        default=os.getenv("OUTPUT_FOLDER", "evaluation_results"),
        help="Output folder (default: evaluation_results or OUTPUT_FOLDER)",
    )
    parser.add_argument(
        "--poppler",
        default=os.getenv("POPPLER_PATH"),
        help="Poppler bin path if you set USE_POPPLER_FOR_PDF (or POPPLER_PATH env)",
    )
    args = parser.parse_args()

    input_path = (args.input_file or os.getenv("INPUT_FILE_PATH") or "").strip()
    if not input_path:
        parser.error(
            "Missing input path. Pass the file as an argument, e.g.:\n"
            "  python -m gt_generator_llms.gt_generator_gateway path\\to\\doc.pdf\n"
            "Or set environment variable INPUT_FILE_PATH."
        )
    if not os.path.isfile(input_path):
        parser.error(f"File not found: {input_path}")

    process_document(input_path, args.output, args.poppler)


if __name__ == "__main__":
    main()
