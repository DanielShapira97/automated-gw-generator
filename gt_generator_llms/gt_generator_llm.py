import os
import base64
import docx
import fitz  # PyMuPDF
from dotenv import load_dotenv
from pdf2image import convert_from_path
from openai import OpenAI

# Load configuration
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_images_from_pdf(pdf_path, doc_folder):
    """Extracts embedded images from PDF and saves them in the document folder."""
    doc = fitz.open(pdf_path)
    image_count = 0
    for page_index in range(len(doc)):
        for img_index, img in enumerate(doc.get_page_images(page_index)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            image_count += 1
            img_filename = f"image_{image_count}.{image_ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(image_bytes)
    print(f"Extracted {image_count} images from PDF.")

def extract_images_from_docx(docx_path, doc_folder):
    """Extracts images from Word document and saves them in the document folder."""
    doc = docx.Document(docx_path)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            img_data = rel.target_part.blob
            ext = rel.target_ref.split('.')[-1]
            img_filename = f"image_{image_count}.{ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(img_data)
    print(f"Extracted {image_count} images from DOCX.")

def extract_raw_text(input_path, doc_folder, poppler_path):
    """Extracts raw text and handles image extraction based on file type."""
    file_ext = os.path.splitext(input_path)[1].lower()
    full_raw_content = ""

    if file_ext == ".docx":
        doc = docx.Document(input_path)
        for para in doc.paragraphs:
            if para.text.strip(): full_raw_content += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                full_raw_content += " | ".join([cell.text.strip() for cell in row.cells]) + "\n"
        extract_images_from_docx(input_path, doc_folder)
        
    elif file_ext == ".pdf":
        extract_images_from_pdf(input_path, doc_folder)
        pages = convert_from_path(input_path, poppler_path=poppler_path)
        for i, page in enumerate(pages):
            print(f"Vision Processing: Page {i+1}...")
            temp_path = os.path.join(doc_folder, f"temp_p{i}.jpg")
            page.save(temp_path, "JPEG")
            
            with open(temp_path, "rb") as f:
                base64_img = base64.b64encode(f.read()).decode('utf-8')
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": "Transcribe ALL text from this image exactly. No data loss. Format any tables using | and - characters to draw them. Do NOT extract text from pictures, diagrams, or charts. Do NOT describe images or pictures. No markdown formatting other than for tables."},
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_img}"}}
                ]}]
            )
            full_raw_content += response.choices[0].message.content + "\n\n"
            os.remove(temp_path) # Cleanup temporary page snapshots
            
    return full_raw_content

def structure_into_blocks(raw_text):
    """Groups text into large blocks with zero content removal."""
    prompt = """
    Divide the text into LARGE logical blocks using ==== markers.
    1. DO NOT DELETE ANY TEXT. NO SUMMARIZING.
    2. NO MARKDOWN symbols except for tables using | and -.
    3. Every block must start and end with ====.
    4. Return ONLY the blocks.
    """
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": f"{prompt}\n\nRAW TEXT:\n{raw_text}"}]
    )
    return response.choices[0].message.content.strip()

def main():
    input_path = os.getenv("INPUT_FILE_PATH") or os.getenv("INPUT_PDF_PATH")
    base_output_dir = os.getenv("OUTPUT_FOLDER", "evaluation_results")
    poppler_path = os.getenv("POPPLER_PATH")

    if not input_path or not os.path.exists(input_path):
        print(f"Error: File {input_path} not found.")
        return

    # Create document-specific folder
    file_name_full = os.path.basename(input_path)
    doc_name = os.path.splitext(file_name_full)[0]
    doc_folder = os.path.join(base_output_dir, doc_name)
    
    if not os.path.exists(doc_folder):
        os.makedirs(doc_folder)

    # 1. Extract raw text and specific images
    raw_content = extract_raw_text(input_path, doc_folder, poppler_path)
    
    # 2. Structure into blocks
    final_gt = structure_into_blocks(raw_content)
    
    # 3. Save final Ground Truth file
    gt_file_name = f"{doc_name} ground truth.txt"
    result_path = os.path.join(doc_folder, gt_file_name)
    
    with open(result_path, "w", encoding="utf-8") as f:
        f.write(final_gt)
    
    print(f"\n--- SUCCESS ---")
    print(f"Folder: {doc_folder}")
    print(f"File: {gt_file_name}")

if __name__ == "__main__":
    main()