import os
import docx
import fitz
import google.generativeai as genai
from dotenv import load_dotenv
from pdf2image import convert_from_path
import PIL.Image

# Load configuration
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

# UPDATED: Using Gemini 3 Flash Preview as per your available models list
MODEL_NAME = 'models/gemini-3-flash-preview' 
model = genai.GenerativeModel(MODEL_NAME)

def extract_images_from_pdf(pdf_path, doc_folder):
    """Extracts embedded images from PDF and saves them in the document folder."""
    doc = fitz.open(pdf_path)
    image_count = 0
    for page_index in range(len(doc)):
        for img_index, img in enumerate(doc.get_page_images(page_index)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]
            
            image_count += 1
            img_filename = f"image_{image_count}.{image_ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(image_bytes)
    print(f"Extracted {image_count} images from PDF.")

def extract_images_from_docx(docx_path, doc_folder):
    """Extracts images from Word document and saves them in the document folder."""
    doc = docx.Document(docx_path)
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            img_data = rel.target_part.blob
            ext = rel.target_ref.split('.')[-1]
            img_filename = f"image_{image_count}.{ext}"
            with open(os.path.join(doc_folder, img_filename), "wb") as f:
                f.write(img_data)
    print(f"Extracted {image_count} images from DOCX.")

def get_raw_text_content(input_path, doc_folder, poppler_path):
    file_ext = os.path.splitext(input_path)[1].lower()
    raw_content = ""

    if file_ext == ".docx":
        print(f"Reading DOCX: {input_path}")
        doc = docx.Document(input_path)
        for p in doc.paragraphs: raw_content += p.text + "\n"
        for t in doc.tables:
            for r in t.rows: raw_content += " | ".join([c.text.strip() for c in r.cells]) + "\n"
        extract_images_from_docx(input_path, doc_folder)
        
    elif file_ext == ".pdf":
        extract_images_from_pdf(input_path, doc_folder)
        doc = fitz.open(input_path)
        print(f"Total pages to process: {len(doc)}")
        
        for i in range(len(doc)):
            print(f"--- Processing Page {i+1} ---")
            page = doc.load_page(i)
            # Render page to an image (pixmap)
            pix = page.get_pixmap(dpi=200) 
            temp_path = os.path.join(doc_folder, f"p_{i}.jpg")
            pix.save(temp_path)
            
            with PIL.Image.open(temp_path) as img:
                try:
                    # Direct extraction without sleep
                    response = model.generate_content(["Extract all text from this image exactly. No data loss. Format any tables using | and - characters to draw them. Do NOT extract text from pictures, diagrams, or charts. Do NOT describe images or pictures. No markdown formatting other than for tables.", img])
                    if response.text:
                        raw_content += response.text + "\n\n"
                except Exception as e:
                    print(f"Error on Page {i+1}: {e}")
            
            os.remove(temp_path)
            
    return raw_content

def structure_into_semantic_blocks(raw_text):
    if not raw_text.strip(): return "Error: No text extracted."
    print("Organizing text into logical blocks...")
    prompt = f"Divide the text into LARGE logical blocks using ==== markers. 1. DO NOT DELETE ANY TEXT. NO SUMMARIZING. 2. NO MARKDOWN symbols except for tables using | and -. 3. Every block must start and end with ====. 4. Return ONLY the blocks. Text:\n\n{raw_text}"
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"Structuring error: {e}")
        return raw_text

def process_document(input_path, output_base="evaluation_results", poppler=None):
    if not input_path:
        print("Missing input_path")
        return None

    doc_name = os.path.splitext(os.path.basename(input_path))[0]
    doc_folder = os.path.join(output_base, doc_name)
    os.makedirs(doc_folder, exist_ok=True)

    # 1. Extraction (No Sleep)
    raw_text = get_raw_text_content(input_path, doc_folder, poppler)
    
    # 2. Blocks
    final_text = structure_into_semantic_blocks(raw_text)

    # 3. Save
    out_file = os.path.join(doc_folder, f"{doc_name} ground truth.txt")
    with open(out_file, "w", encoding="utf-8") as f:
        f.write(final_text)
    
    print(f"\n--- SUCCESS! Created: {out_file} ---")
    return out_file

def main():
    input_path = os.getenv("INPUT_FILE_PATH")
    output_base = os.getenv("OUTPUT_FOLDER", "evaluation_results")
    poppler = os.getenv("POPPLER_PATH")
    process_document(input_path, output_base, poppler)

if __name__ == "__main__":
    main()