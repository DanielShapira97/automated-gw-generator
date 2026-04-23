import os
import docx
import fitz
import google.generativeai as genai
from dotenv import load_dotenv
from pdf2image import convert_from_path
import PIL.Image

# Load configuration
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=api_key)

# UPDATED: Using Gemini 3 Flash Preview as per your available models list
MODEL_NAME = 'models/gemini-3-flash-preview' 
model = genai.GenerativeModel(MODEL_NAME)

def get_raw_text_content(input_path, doc_folder, poppler_path):
    file_ext = os.path.splitext(input_path)[1].lower()
    raw_content = ""

    if file_ext == ".docx":
        print(f"Reading DOCX: {input_path}")
        doc = docx.Document(input_path)
        for p in doc.paragraphs: raw_content += p.text + "\n"
        for t in doc.tables:
            for r in t.rows: raw_content += " | ".join([c.text.strip() for c in r.cells]) + "\n"
        
    elif file_ext == ".pdf":
        pages = convert_from_path(input_path, poppler_path=poppler_path)
        print(f"Total pages to process: {len(pages)}")
        
        for i, page in enumerate(pages):
            print(f"--- Processing Page {i+1} ---")
            temp_path = os.path.join(doc_folder, f"p_{i}.jpg")
            page.save(temp_path, "JPEG")
            img = PIL.Image.open(temp_path)
            
            try:
                # Direct extraction without sleep
                response = model.generate_content(["Extract all text from this image exactly. No data loss.NO MARKDOWN symbols", img])
                if response.text:
                    raw_content += response.text + "\n\n"
            except Exception as e:
                print(f"Error on Page {i+1}: {e}")
            
            os.remove(temp_path)
            
    return raw_content

def structure_into_semantic_blocks(raw_text):
    if not raw_text.strip(): return "Error: No text extracted."
    print("Organizing text into logical blocks...")
    prompt = f"Divide the text into LARGE logical blocks using ==== markers. 1. DO NOT DELETE ANY TEXT. NO SUMMARIZING. 2. NO MARKDOWN symbols!!!! Plain text only. 3. Every block must start and end with ====. 4. Return ONLY the blocks. Text:\n\n{raw_text}"
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"Structuring error: {e}")
        return raw_text

def main():
    input_path = os.getenv("INPUT_FILE_PATH")
    output_base = os.getenv("OUTPUT_FOLDER", "evaluation_results")
    poppler = os.getenv("POPPLER_PATH")

    if not input_path:
        print("Missing INPUT_FILE_PATH in .env")
        return

    doc_name = os.path.splitext(os.path.basename(input_path))[0]
    doc_folder = os.path.join(output_base, doc_name)
    os.makedirs(doc_folder, exist_ok=True)

    # 1. Extraction (No Sleep)
    raw_text = get_raw_text_content(input_path, doc_folder, poppler)
    
    # 2. Blocks
    final_text = structure_into_semantic_blocks(raw_text)

    # 3. Save
    out_file = os.path.join(doc_folder, f"{doc_name} ground truth.txt")
    with open(out_file, "w", encoding="utf-8") as f:
        f.write(final_text)
    
    print(f"\n--- SUCCESS! Created: {out_file} ---")

if __name__ == "__main__":
    main()